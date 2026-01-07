"""
IPO-Compliant Patent Document Export with Diagrams
===================================================
Generates professional PDF and DOCX exports in exact Indian Patent Office format.
Includes all diagrams/figures embedded properly.

IPO Form 2 Complete Structure:
1. Title Page (THE PATENTS ACT, 1970)
2. Abstract
3. Title of Invention
4. Field of the Invention
5. Background of the Invention
6. Objects of the Invention
7. Summary of the Invention
8. Brief Description of the Drawings
9. Detailed Description (with reference numerals)
10. Industrial Applicability
11. Claims (with "WE CLAIM:")
12. Abstract (repeat for filing)
13. Drawing Sheets (Figures 1, 2, 3...)
"""

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image as RLImage
from reportlab.lib import colors
from datetime import datetime
from typing import Dict, Optional, List
import os


class IPOPatentPDFGenerator:
    """Generate IPO-compliant patent PDF documents with diagrams."""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_styles()
    
    def _setup_styles(self):
        """Setup IPO-compliant text styles."""
        self.styles.add(ParagraphStyle(
            name='PatentTitle',
            fontSize=14,
            leading=18,
            alignment=TA_CENTER,
            spaceAfter=20,
            spaceBefore=20,
            fontName='Times-Bold'
        ))
        
        self.styles.add(ParagraphStyle(
            name='SectionHeader',
            fontSize=12,
            leading=16,
            alignment=TA_LEFT,
            spaceAfter=12,
            spaceBefore=20,
            fontName='Times-Bold'
        ))
        
        self.styles.add(ParagraphStyle(
            name='PatentBody',
            fontSize=11,
            leading=16,
            alignment=TA_JUSTIFY,
            spaceAfter=8,
            fontName='Times-Roman',
            firstLineIndent=0.5*inch
        ))
        
        self.styles.add(ParagraphStyle(
            name='ClaimText',
            fontSize=11,
            leading=16,
            alignment=TA_JUSTIFY,
            spaceAfter=12,
            fontName='Times-Roman',
            leftIndent=0.25*inch
        ))
        
        self.styles.add(ParagraphStyle(
            name='AbstractText',
            fontSize=11,
            leading=16,
            alignment=TA_JUSTIFY,
            spaceAfter=8,
            fontName='Times-Roman'
        ))
        
        self.styles.add(ParagraphStyle(
            name='PageHeader',
            fontSize=10,
            leading=12,
            alignment=TA_CENTER,
            fontName='Times-Roman'
        ))
        
        self.styles.add(ParagraphStyle(
            name='FigureCaption',
            fontSize=10,
            leading=14,
            alignment=TA_CENTER,
            spaceAfter=20,
            spaceBefore=10,
            fontName='Times-Roman'
        ))
    
    def create_patent_pdf(
        self,
        sections: Dict[str, str],
        output_path: str = "patent_ipo_format.pdf",
        applicant_name: str = "",
        inventor_name: str = "",
        diagram_paths: List[str] = None
    ) -> str:
        """
        Create IPO-compliant patent PDF with diagrams.
        
        Args:
            sections: Dict with all patent sections
            output_path: Output file path
            applicant_name: Applicant name
            inventor_name: Inventor name
            diagram_paths: List of paths to diagram images
        
        Returns:
            Path to generated PDF
        """
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=1*inch,
            leftMargin=1*inch,
            topMargin=1*inch,
            bottomMargin=1*inch
        )
        
        story = []
        diagram_paths = diagram_paths or []
        
        # ========== PAGE 1: TITLE PAGE ==========
        story.append(Paragraph("THE PATENTS ACT, 1970", self.styles['PageHeader']))
        story.append(Paragraph("(39 of 1970)", self.styles['PageHeader']))
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph("COMPLETE SPECIFICATION", self.styles['PatentTitle']))
        story.append(Paragraph("(See Section 10; Rule 13)", self.styles['PageHeader']))
        story.append(Spacer(1, 0.5*inch))
        
        # Title
        title = sections.get("Title", "[TITLE NOT PROVIDED]")
        story.append(Paragraph("<b>TITLE OF THE INVENTION</b>", self.styles['SectionHeader']))
        story.append(Paragraph(title.upper(), self.styles['PatentTitle']))
        story.append(Spacer(1, 0.3*inch))
        
        # Applicant/Inventor Info
        if applicant_name:
            story.append(Paragraph(f"<b>Applicant:</b> {applicant_name}", self.styles['PatentBody']))
        if inventor_name:
            story.append(Paragraph(f"<b>Inventor:</b> {inventor_name}", self.styles['PatentBody']))
        
        story.append(PageBreak())
        
        # ========== ABSTRACT PAGE ==========
        story.append(Paragraph("<b>ABSTRACT</b>", self.styles['SectionHeader']))
        abstract = sections.get("Abstract", "[ABSTRACT NOT PROVIDED]")
        story.append(Paragraph(abstract, self.styles['AbstractText']))
        story.append(PageBreak())
        
        # ========== SPECIFICATION SECTIONS ==========
        section_order = [
            ("FIELD OF THE INVENTION", "Field of the Invention"),
            ("BACKGROUND OF THE INVENTION", "Background of the Invention"),
            ("OBJECTS OF THE INVENTION", "Objects of the Invention"),
            ("SUMMARY OF THE INVENTION", "Summary of the Invention"),
            ("BRIEF DESCRIPTION OF THE DRAWINGS", "Brief Description of the Drawings"),
            ("DETAILED DESCRIPTION OF THE INVENTION", "Detailed Description of the Invention"),
            ("INDUSTRIAL APPLICABILITY", "Industrial Applicability"),
        ]
        
        for header_text, section_key in section_order:
            content = sections.get(section_key, "")
            if not content or content == "[Not Generated]":
                continue
            
            story.append(Paragraph(f"<b>{header_text}</b>", self.styles['SectionHeader']))
            
            # Split content into paragraphs
            paragraphs = content.strip().split('\n\n')
            for para in paragraphs:
                if para.strip():
                    clean_para = para.strip().replace('\n', ' ')
                    story.append(Paragraph(clean_para, self.styles['PatentBody']))
            
            story.append(Spacer(1, 0.2*inch))
        
        # ========== CLAIMS (New Page) ==========
        story.append(PageBreak())
        story.append(Paragraph("<b>CLAIMS</b>", self.styles['SectionHeader']))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph("<b>WE CLAIM:</b>", self.styles['PatentBody']))
        story.append(Spacer(1, 0.2*inch))
        
        claims = sections.get("Claims", "[CLAIMS NOT PROVIDED]")
        claim_lines = claims.strip().split('\n')
        for line in claim_lines:
            if line.strip():
                story.append(Paragraph(line.strip(), self.styles['ClaimText']))
        
        # ========== DRAWING SHEETS (New Pages) ==========
        if diagram_paths:
            story.append(PageBreak())
            story.append(Paragraph("<b>DRAWING SHEETS</b>", self.styles['SectionHeader']))
            story.append(Spacer(1, 0.3*inch))
            
            fig_num = 1
            for diagram_path in diagram_paths:
                if os.path.exists(diagram_path):
                    try:
                        # Add figure with proper sizing
                        img = RLImage(diagram_path, width=5*inch, height=4*inch)
                        story.append(img)
                        
                        # Figure caption
                        fig_name = os.path.basename(diagram_path).replace('.png', '').replace('_', ' ').title()
                        story.append(Paragraph(f"<b>Figure {fig_num}:</b> {fig_name}", self.styles['FigureCaption']))
                        story.append(Spacer(1, 0.3*inch))
                        
                        fig_num += 1
                    except Exception as e:
                        story.append(Paragraph(f"[Figure {fig_num} - Error loading image]", self.styles['FigureCaption']))
                        fig_num += 1
        else:
            # No diagrams provided - check Brief Description for figure count and add placeholders
            brief_desc = sections.get("Brief Description of the Drawings", "")
            if brief_desc:
                import re
                figure_matches = re.findall(r'Figure\s+(\d+)', brief_desc, re.IGNORECASE)
                if figure_matches:
                    story.append(PageBreak())
                    story.append(Paragraph("<b>DRAWING SHEETS</b>", self.styles['SectionHeader']))
                    story.append(Spacer(1, 0.3*inch))
                    story.append(Paragraph("<i>Note: Please insert drawings below. Placeholders are provided based on your Brief Description of the Drawings.</i>", self.styles['PatentBody']))
                    story.append(Spacer(1, 0.3*inch))
                    
                    max_fig = max(int(m) for m in figure_matches)
                    for fig_num in range(1, max_fig + 1):
                        # Create placeholder box
                        story.append(Spacer(1, 0.5*inch))
                        story.append(Paragraph(f"<b>[INSERT FIGURE {fig_num} DRAWING HERE]</b>", self.styles['SectionHeader']))
                        story.append(Spacer(1, 2*inch))  # Space for drawing
                        story.append(Paragraph(f"<b>Figure {fig_num}</b>", self.styles['FigureCaption']))
                        story.append(Spacer(1, 0.3*inch))
        
        # ========== FOOTER ==========
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph("---", self.styles['PageHeader']))
        
        dated = datetime.now().strftime("%d %B %Y")
        story.append(Paragraph(f"Dated this {dated}", self.styles['PageHeader']))
        
        if applicant_name:
            story.append(Spacer(1, 0.3*inch))
            story.append(Paragraph(f"Signature: _______________________", self.styles['PageHeader']))
            story.append(Paragraph(f"Applicant: {applicant_name}", self.styles['PageHeader']))
        
        # Build PDF
        doc.build(story)
        return output_path


def create_patent_pdf(sections: Dict[str, str], output_path: str = "generated_patent.pdf", diagram_paths: List[str] = None) -> str:
    """Convenience function to create IPO-compliant PDF."""
    generator = IPOPatentPDFGenerator()
    return generator.create_patent_pdf(sections, output_path, diagram_paths=diagram_paths)


def create_patent_docx(
    sections: Dict[str, str],
    output_path: str = "patent_ipo_format.docx",
    applicant_name: str = "",
    inventor_name: str = "",
    diagram_paths: List[str] = None
) -> str:
    """
    Create IPO-compliant patent DOCX document with diagrams.
    """
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    diagram_paths = diagram_paths or []
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    
    # ========== TITLE PAGE ==========
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run("THE PATENTS ACT, 1970\n(39 of 1970)")
    header_run.font.name = 'Times New Roman'
    header_run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    spec_para = doc.add_paragraph()
    spec_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spec_run = spec_para.add_run("COMPLETE SPECIFICATION")
    spec_run.bold = True
    spec_run.font.name = 'Times New Roman'
    spec_run.font.size = Pt(14)
    
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("(See Section 10; Rule 13)")
    sub_run.font.name = 'Times New Roman'
    sub_run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Title
    title_header = doc.add_paragraph()
    title_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    th_run = title_header.add_run("TITLE OF THE INVENTION")
    th_run.bold = True
    th_run.font.name = 'Times New Roman'
    th_run.font.size = Pt(12)
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(sections.get("Title", "[TITLE]").upper())
    title_run.bold = True
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(14)
    
    if applicant_name:
        app_info = doc.add_paragraph()
        app_run = app_info.add_run(f"Applicant: {applicant_name}")
        app_run.font.name = 'Times New Roman'
        app_run.font.size = Pt(11)
    
    if inventor_name:
        inv_info = doc.add_paragraph()
        inv_run = inv_info.add_run(f"Inventor: {inventor_name}")
        inv_run.font.name = 'Times New Roman'
        inv_run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # ========== ABSTRACT ==========
    abs_header = doc.add_paragraph()
    abs_run = abs_header.add_run("ABSTRACT")
    abs_run.bold = True
    abs_run.font.name = 'Times New Roman'
    abs_run.font.size = Pt(12)
    
    abs_para = doc.add_paragraph()
    abs_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abs_content = abs_para.add_run(sections.get("Abstract", "[ABSTRACT]"))
    abs_content.font.name = 'Times New Roman'
    abs_content.font.size = Pt(11)
    
    doc.add_page_break()
    
    # ========== ALL SPECIFICATION SECTIONS ==========
    section_order = [
        ("FIELD OF THE INVENTION", "Field of the Invention"),
        ("BACKGROUND OF THE INVENTION", "Background of the Invention"),
        ("OBJECTS OF THE INVENTION", "Objects of the Invention"),
        ("SUMMARY OF THE INVENTION", "Summary of the Invention"),
        ("BRIEF DESCRIPTION OF THE DRAWINGS", "Brief Description of the Drawings"),
        ("DETAILED DESCRIPTION OF THE INVENTION", "Detailed Description of the Invention"),
        ("INDUSTRIAL APPLICABILITY", "Industrial Applicability"),
    ]
    
    for header_text, section_key in section_order:
        content = sections.get(section_key, "")
        if not content or content == "[Not Generated]":
            continue
        
        sec_header = doc.add_paragraph()
        sec_run = sec_header.add_run(header_text)
        sec_run.bold = True
        sec_run.font.name = 'Times New Roman'
        sec_run.font.size = Pt(12)
        
        sec_para = doc.add_paragraph()
        sec_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        sec_content = sec_para.add_run(content.strip())
        sec_content.font.name = 'Times New Roman'
        sec_content.font.size = Pt(11)
        
        doc.add_paragraph()
    
    # ========== CLAIMS ==========
    doc.add_page_break()
    
    claims_header = doc.add_paragraph()
    claims_run = claims_header.add_run("CLAIMS")
    claims_run.bold = True
    claims_run.font.name = 'Times New Roman'
    claims_run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    we_claim = doc.add_paragraph()
    wc_run = we_claim.add_run("WE CLAIM:")
    wc_run.bold = True
    wc_run.font.name = 'Times New Roman'
    wc_run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    claims_content = sections.get("Claims", "[CLAIMS]")
    for line in claims_content.strip().split('\n'):
        if line.strip():
            claim_para = doc.add_paragraph()
            claim_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            claim_run = claim_para.add_run(line.strip())
            claim_run.font.name = 'Times New Roman'
            claim_run.font.size = Pt(11)
    
    # ========== DRAWING SHEETS ==========
    if diagram_paths:
        doc.add_page_break()
        
        drawing_header = doc.add_paragraph()
        dh_run = drawing_header.add_run("DRAWING SHEETS")
        dh_run.bold = True
        dh_run.font.name = 'Times New Roman'
        dh_run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        fig_num = 1
        for diagram_path in diagram_paths:
            if os.path.exists(diagram_path):
                try:
                    doc.add_picture(diagram_path, width=Inches(5))
                    
                    # Caption
                    caption = doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    fig_name = os.path.basename(diagram_path).replace('.png', '').replace('_', ' ').title()
                    cap_run = caption.add_run(f"Figure {fig_num}: {fig_name}")
                    cap_run.font.name = 'Times New Roman'
                    cap_run.font.size = Pt(10)
                    cap_run.bold = True
                    
                    doc.add_paragraph()
                    fig_num += 1
                except Exception as e:
                    error_para = doc.add_paragraph()
                    error_para.add_run(f"[Figure {fig_num} - Image not available]")
                    fig_num += 1
    
    # ========== FOOTER ==========
    doc.add_paragraph()
    doc.add_paragraph()
    
    dated = datetime.now().strftime("%d %B %Y")
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Dated this {dated}")
    date_run.font.name = 'Times New Roman'
    date_run.font.size = Pt(10)
    
    if applicant_name:
        sig_para = doc.add_paragraph()
        sig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sig_run = sig_para.add_run("\n\nSignature: _______________________")
        sig_run.font.name = 'Times New Roman'
        sig_run.font.size = Pt(10)
        
        app_para = doc.add_paragraph()
        app_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        app_run = app_para.add_run(f"Applicant: {applicant_name}")
        app_run.font.name = 'Times New Roman'
        app_run.font.size = Pt(10)
    
    doc.save(output_path)
    return output_path


def get_diagram_paths():
    """Get paths to generated diagrams."""
    diagram_dir = os.path.join(os.path.dirname(__file__), "generated_diagrams")
    paths = []
    
    if os.path.exists(diagram_dir):
        for f in sorted(os.listdir(diagram_dir)):
            if f.endswith('.png'):
                paths.append(os.path.join(diagram_dir, f))
    
    return paths


# Test
if __name__ == "__main__":
    test_sections = {
        "Title": "A Smart Industrial Monitoring System",
        "Abstract": "The present invention relates to a smart monitoring system for industrial environments comprising sensor units (101), a central processing hub (102), and an alert module (103) for real-time anomaly detection and notification.",
        "Field of the Invention": "The present invention relates to the field of industrial automation and monitoring systems, more particularly to an intelligent system for real-time monitoring and anomaly detection in industrial environments.",
        "Background of the Invention": "Conventional industrial monitoring systems rely on manual inspection and basic sensor readings. These systems lack intelligent analysis capabilities and fail to provide real-time anomaly detection. There exists a need for an automated, intelligent monitoring solution that can process sensor data in real-time and generate immediate alerts.",
        "Objects of the Invention": "The primary object of the present invention is to provide an intelligent industrial monitoring system.\n\nAnother object is to provide real-time anomaly detection capabilities.\n\nYet another object is to minimize industrial accidents through predictive maintenance.",
        "Summary of the Invention": "According to one aspect of the invention, there is provided a smart industrial monitoring system comprising: a plurality of sensor units (101) configured to collect environmental data; a central processing hub (102) connected to said sensor units; and an alert module (103) for generating notifications.",
        "Brief Description of the Drawings": "Figure 1 shows a block diagram of the smart monitoring system (100) according to the present invention.\n\nFigure 2 shows a flowchart illustrating the method of operation.\n\nFigure 3 shows a sequence diagram of component interactions.",
        "Detailed Description of the Invention": "Referring now to Figure 1, the smart industrial monitoring system (100) comprises a sensor unit (101), a processing hub (102), and an alert module (103). The sensor unit (101) includes temperature sensors, pressure sensors, and vibration sensors. The processing hub (102) receives data from the sensor unit (101) and performs intelligent analysis. The alert module (103) generates notifications when anomalies are detected.",
        "Industrial Applicability": "The present invention finds application in various industries including manufacturing, oil and gas, chemical processing, and power generation. The system enables predictive maintenance and reduces downtime.",
        "Claims": "1. A smart industrial monitoring system comprising:\n   a plurality of sensor units (101);\n   a central processing hub (102) connected to said sensor units; and\n   an alert module (103).\n\n2. The system as claimed in claim 1, wherein the sensor units include temperature, pressure, and vibration sensors.\n\n3. The system as claimed in claim 1, wherein the processing hub performs real-time anomaly detection.\n\n4. A method for industrial monitoring comprising:\n   collecting sensor data;\n   analyzing said data; and\n   generating alerts."
    }
    
    print("=" * 60)
    print("IPO PATENT EXPORT WITH DIAGRAMS - TEST")
    print("=" * 60)
    
    # Get diagram paths
    diagrams = get_diagram_paths()
    print(f"Found {len(diagrams)} diagrams: {diagrams}")
    
    # Test PDF
    pdf_path = create_patent_pdf(test_sections, "test_ipo_complete.pdf", diagram_paths=diagrams)
    print(f"✅ PDF created: {pdf_path}")
    
    # Test DOCX
    docx_path = create_patent_docx(test_sections, "test_ipo_complete.docx", 
                                   applicant_name="Test Applicant",
                                   inventor_name="Test Inventor",
                                   diagram_paths=diagrams)
    print(f"✅ DOCX created: {docx_path}")
    
    print()
    print("IPO SECTIONS INCLUDED:")
    print("  1. Title Page (THE PATENTS ACT, 1970)")
    print("  2. Abstract")
    print("  3. Field of the Invention")
    print("  4. Background of the Invention")
    print("  5. Objects of the Invention")
    print("  6. Summary of the Invention")
    print("  7. Brief Description of the Drawings")
    print("  8. Detailed Description")
    print("  9. Industrial Applicability")
    print("  10. Claims (with WE CLAIM:)")
    print("  11. Drawing Sheets (Figures)")
    print("  12. Signature Block")
